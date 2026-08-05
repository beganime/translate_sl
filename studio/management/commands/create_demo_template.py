from pathlib import Path
from tempfile import TemporaryDirectory

from django.core.files import File
from django.core.management.base import BaseCommand
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from studio.models import DocumentTemplate, TemplateVariable


class Command(BaseCommand):
    help = "Создать демонстрационный DOCX-шаблон и набор Jinja-переменных"

    def handle(self, *args, **options):
        template, created = DocumentTemplate.objects.get_or_create(
            name="Демонстрационный перевод",
            defaults={"description": "Простой шаблон для проверки полного цикла."},
        )
        if created or not template.file:
            with TemporaryDirectory() as tmp:
                path = Path(tmp) / "demo-template.docx"
                doc = Document()
                section = doc.sections[0]
                section.top_margin = Mm(22)
                section.bottom_margin = Mm(22)
                section.left_margin = Mm(25)
                section.right_margin = Mm(25)
                normal = doc.styles["Normal"]
                normal.font.name = "Arial"
                normal.font.size = Pt(11)
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title.add_run("ПЕРЕВОД ДОКУМЕНТА")
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(16)
                doc.add_paragraph()
                for label, variable in [
                    ("Тип документа", "document_type"),
                    ("Номер", "document_number"),
                    ("Дата выдачи", "issue_date"),
                    ("Фамилия, имя", "full_name"),
                    ("Содержание", "body_text"),
                    ("Орган выдачи", "issuer"),
                ]:
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(f"{label}: ").bold = True
                    paragraph.add_run("{{ " + variable + " }}")
                doc.save(path)
                with path.open("rb") as stream:
                    template.file.save(path.name, File(stream), save=True)

        variables = [
            ("document_type", "Тип документа", "Определи официальное название документа."),
            ("document_number", "Номер документа", "Найди серию и номер, сохрани исходные символы."),
            ("issue_date", "Дата выдачи", "Приведи дату в формате ДД.ММ.ГГГГ, если это возможно."),
            ("full_name", "Фамилия, имя", "Полное имя держателя документа."),
            ("body_text", "Основной текст", "Переведи основной смысловой текст полностью."),
            ("issuer", "Орган выдачи", "Организация или орган, выдавший документ."),
        ]
        for order, (name, label, instruction) in enumerate(variables, 1):
            TemplateVariable.objects.update_or_create(
                template=template,
                name=name,
                defaults={"label": label, "ai_instruction": instruction, "sort_order": order * 10},
            )
        self.stdout.write(self.style.SUCCESS("Демонстрационный шаблон создан."))

