import tempfile
from pathlib import Path
from unittest.mock import patch

from django.core.files import File
from django.core.files.uploadedfile import SimpleUploadedFile
from django.contrib.auth import get_user_model
from django.core import signing
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document

from .models import AIConfiguration, DocumentTemplate, SourceDocument, TemplateVariable
from .forms import UploadDocumentForm
from .services.gemini import _response_schema, _source_parts
from .services.rendering import render_document


class ModelTests(TestCase):
    def test_api_key_round_trip_is_encrypted(self):
        config = AIConfiguration(name="Gemini")
        config.set_api_key("secret-value")
        config.save()
        self.assertNotIn("secret-value", config.api_key_encrypted)
        self.assertEqual(config.get_api_key(), "secret-value")

    def test_only_one_configuration_stays_active(self):
        first = AIConfiguration.objects.create(name="One", is_active=True)
        AIConfiguration.objects.create(name="Two", is_active=True)
        first.refresh_from_db()
        self.assertFalse(first.is_active)

    def test_schema_uses_template_instructions(self):
        template = DocumentTemplate.objects.create(name="T", file="templates/t.docx")
        variable = TemplateVariable.objects.create(
            template=template, name="full_name", label="Имя", ai_instruction="Найди имя", required=True
        )
        schema = _response_schema([variable])
        self.assertIn("full_name", schema["properties"])
        self.assertIn("Найди имя", schema["properties"]["full_name"]["description"])
        self.assertNotIn("additionalProperties", schema)

    def test_upload_form_accepts_supported_business_formats(self):
        for name in ("scan.pdf", "scan.docx", "scan.jpg", "scan.png"):
            upload = SimpleUploadedFile(name, b"test")
            form = UploadDocumentForm(files={"source_pdf": upload})
            self.assertNotIn("source_pdf", form.errors)

    def test_docx_source_is_extracted_for_gemini(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "source.docx"
            word = Document()
            word.add_paragraph("Passport number A123")
            word.save(source)
            with source.open("rb") as stream:
                stored = SimpleUploadedFile("source.docx", stream.read())
            parts = _source_parts(stored)
            self.assertIn("Passport number A123", parts[0]["text"])


class ViewTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(username="manager", password="safe-test-password")
        self.template = DocumentTemplate.objects.create(name="Свидетельство", file="templates/t.docx")

    def test_anonymous_user_is_redirected_to_login(self):
        response = self.client.get(reverse("studio:dashboard"))
        self.assertRedirects(response, f"{reverse('login')}?next={reverse('studio:dashboard')}")

    def test_dashboard_opens(self):
        self.client.force_login(self.user)
        response = self.client.get(reverse("studio:dashboard"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Скан становится")

    def test_health_checks_database(self):
        response = self.client.get(reverse("health"))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "ok")

    @override_settings(MANAGER_SL_SSO_SECRET="shared-test-secret", MANAGER_SL_SSO_MAX_AGE=120)
    def test_manager_sl_sso_creates_local_session_and_keeps_client_context(self):
        token = signing.dumps(
            {
                "email": "manager@manager-sl.ru",
                "first_name": "Олеся",
                "last_name": "Менеджер",
                "is_staff": False,
                "next": "/upload/?client=SL-2027-001",
            },
            key="shared-test-secret",
            salt="manager-sl.translate-sso.v1",
            compress=True,
        )

        response = self.client.get(reverse("manager_sl_sso"), {"token": token})

        self.assertRedirects(response, "/upload/?client=SL-2027-001", fetch_redirect_response=False)
        user = get_user_model().objects.get(email="manager@manager-sl.ru")
        self.assertFalse(user.has_usable_password())
        self.assertEqual(user.first_name, "Олеся")
        self.assertEqual(str(self.client.session.get("_auth_user_id")), str(user.pk))

    @override_settings(MANAGER_SL_SSO_SECRET="shared-test-secret", MANAGER_SL_SSO_MAX_AGE=120)
    def test_manager_sl_sso_rejects_invalid_signature(self):
        response = self.client.get(reverse("manager_sl_sso"), {"token": "not-a-token"})
        self.assertEqual(response.status_code, 400)

    def test_dashboard_survives_missing_source_file(self):
        SourceDocument.objects.create(
            title="Missing scan", template=self.template, source_pdf="scans/missing.pdf"
        )
        self.client.force_login(self.user)
        response = self.client.get(reverse("studio:dashboard"))
        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "файл недоступен")

    def test_private_scan_requires_login_and_is_available_after_login(self):
        with tempfile.TemporaryDirectory() as media_dir, override_settings(MEDIA_ROOT=media_dir):
            scan_dir = Path(media_dir) / "scans"
            scan_dir.mkdir()
            (scan_dir / "private.pdf").write_bytes(b"%PDF-1.4\n%%EOF")
            document = SourceDocument.objects.create(
                title="Private", template=self.template, source_pdf="scans/private.pdf"
            )
            url = reverse("studio:file", args=[document.pk, "source_pdf"])
            self.assertEqual(self.client.get(url).status_code, 302)
            self.client.force_login(self.user)
            response = self.client.get(url)
            self.assertEqual(response.status_code, 200)
            self.assertEqual(response["X-Frame-Options"], "SAMEORIGIN")
            response.close()

    @patch("studio.views.render_document")
    def test_regenerate_action_is_preserved(self, mocked_render):
        self.client.force_login(self.user)
        document = SourceDocument.objects.create(
            title="Test", template=self.template, source_pdf="scans/test.pdf", status=SourceDocument.Status.REVIEW
        )
        response = self.client.post(reverse("studio:review", args=[document.pk]), {
            "title": "Updated",
            "action": "regenerate",
        })
        self.assertRedirects(response, reverse("studio:review", args=[document.pk]))
        mocked_render.assert_called_once()

    @patch("studio.views.process_document")
    def test_upload_creates_document_and_starts_processing(self, mocked_process):
        self.client.force_login(self.user)
        pdf = SimpleUploadedFile("scan.pdf", b"%PDF-1.4\n%%EOF", content_type="application/pdf")
        response = self.client.post(reverse("studio:upload"), {
            "title": "Тест",
            "template": self.template.pk,
            "source_pdf": pdf,
        })
        document = SourceDocument.objects.get()
        self.assertRedirects(response, reverse("studio:review", args=[document.pk]))
        mocked_process.assert_called_once()

    @patch("studio.views.process_document")
    @patch("studio.views.download_original", return_value=("passport.pdf", b"%PDF-1.4\n%%EOF"))
    @patch("studio.views.list_originals")
    def test_upload_can_import_original_from_disksl(self, mocked_list, mocked_download, mocked_process):
        disk_key = "disk/2027/Контракт/Иванов Иван (SL-001)/оригиналы/passport.pdf"
        mocked_list.return_value = [{"key": disk_key, "label": "2027/.../passport.pdf", "size": 20}]
        self.client.force_login(self.user)

        response = self.client.post(reverse("studio:upload"), {
            "title": "Паспорт",
            "template": self.template.pk,
            "disk_source": disk_key,
        })

        document = SourceDocument.objects.get()
        self.assertRedirects(response, reverse("studio:review", args=[document.pk]))
        self.assertEqual(document.disk_source_key, disk_key)
        # Django storage is allowed to disambiguate an existing filename
        # (for example ``passport_Ab12Cd.pdf``).  The imported document must
        # retain its safe format and content, not a particular storage suffix.
        self.assertEqual(Path(document.source_pdf.name).suffix.lower(), ".pdf")
        document.source_pdf.open("rb")
        try:
            self.assertEqual(document.source_pdf.read(), b"%PDF-1.4\n%%EOF")
        finally:
            document.source_pdf.close()
        mocked_download.assert_called_once_with(disk_key)
        mocked_process.assert_called_once_with(document)


@override_settings(WORD_PDF_CONVERSION=False)
class RenderingTests(TestCase):
    def test_docx_template_is_rendered(self):
        with tempfile.TemporaryDirectory() as media_dir:
            with override_settings(MEDIA_ROOT=media_dir):
                source_path = Path(media_dir) / "source.docx"
                docx = Document()
                docx.add_paragraph("Имя: {{ full_name }}")
                docx.save(source_path)
                template = DocumentTemplate.objects.create(name="T")
                with source_path.open("rb") as stream:
                    template.file.save("source.docx", File(stream), save=True)
                TemplateVariable.objects.create(template=template, name="full_name", label="Имя")
                document = SourceDocument.objects.create(
                    title="Result", template=template, source_pdf="scans/scan.pdf", edited_data={"full_name": "Иван Иванов"}
                )
                render_document(document)
                document.refresh_from_db()
                self.assertTrue(document.generated_docx.name)
                rendered = Document(document.generated_docx.path)
                self.assertIn("Иван Иванов", "\n".join(p.text for p in rendered.paragraphs))
                user = get_user_model().objects.create_user(username="previewer", password="safe-test-password")
                self.client.force_login(user)
                response = self.client.get(reverse("studio:preview", args=[document.pk]))
                self.assertContains(response, "Иван Иванов")
