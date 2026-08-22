import mimetypes
from pathlib import Path

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.core.files.base import ContentFile
from django.http import FileResponse, Http404, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_POST
from django.views.decorators.clickjacking import xframe_options_sameorigin

from .forms import DocumentReviewForm, UploadDocumentForm
from .models import AIConfiguration, DocumentTemplate, SourceDocument
from .services.rendering import process_document, render_document
from .services.disk import DiskUnavailable, download_original, list_originals, upload_translation


@login_required
def dashboard(request):
    documents = list(SourceDocument.objects.select_related("template")[:50])
    for document in documents:
        try:
            document.source_size = document.source_pdf.size
        except (OSError, ValueError):
            # An old database row may outlive its uploaded file. It should not
            # make the entire manager dashboard unavailable.
            document.source_size = None
    context = {
        "documents": documents,
        "template_count": DocumentTemplate.objects.filter(is_active=True).count(),
        "ready_count": SourceDocument.objects.filter(status=SourceDocument.Status.READY).count(),
        "ai_ready": AIConfiguration.objects.filter(is_active=True).exclude(api_key_encrypted="").exists(),
    }
    return render(request, "studio/dashboard.html", context)


@login_required
def upload_document(request):
    client_sl_id = (request.GET.get("client") or request.POST.get("client_sl_id") or "").strip()[:40]
    try:
        disk_files = list_originals(client_sl_id=client_sl_id or None)
    except Exception:
        disk_files = []
    disk_choices = [(item["key"], item["label"]) for item in disk_files]
    if request.method == "POST":
        form = UploadDocumentForm(request.POST, request.FILES, disk_choices=disk_choices)
        if form.is_valid():
            document = form.save(commit=False)
            disk_source_key = form.cleaned_data.get("disk_source")
            if disk_source_key:
                try:
                    filename, content = download_original(disk_source_key)
                    document.source_pdf.save(filename, ContentFile(content), save=False)
                    document.disk_source_key = disk_source_key
                except (DiskUnavailable, ValueError, OSError) as exc:
                    form.add_error("disk_source", str(exc))
                    return render(request, "studio/upload.html", {
                        "form": form,
                        "disk_files": disk_files,
                        "client_sl_id": client_sl_id,
                    })
            document.save()
            try:
                process_document(document)
                messages.success(request, "Скан распознан. Проверьте найденные значения.")
            except Exception:
                messages.error(request, "Обработка не завершена. Подробности показаны на странице документа.")
            return redirect("studio:review", pk=document.pk)
    else:
        form = UploadDocumentForm(disk_choices=disk_choices)
    return render(request, "studio/upload.html", {
        "form": form,
        "disk_files": disk_files,
        "client_sl_id": client_sl_id,
    })


@login_required
def review_document(request, pk):
    document = get_object_or_404(
        SourceDocument.objects.select_related("template").prefetch_related("template__variables"),
        pk=pk,
    )
    if request.method == "POST":
        form = DocumentReviewForm(request.POST, document=document)
        if form.is_valid():
            document.title = form.cleaned_data["title"]
            document.edited_data = form.variable_values()
            document.error_message = ""
            action = request.POST.get("action", "save")
            try:
                if action in {"regenerate", "finalize"}:
                    render_document(document, document.edited_data)
                if action == "finalize" and document.disk_source_key and document.generated_docx:
                    document.generated_docx.open("rb")
                    try:
                        document.disk_result_key = upload_translation(
                            document.disk_source_key,
                            Path(document.generated_docx.name).name,
                            document.generated_docx.read(),
                        )
                    finally:
                        document.generated_docx.close()
                document.status = document.Status.READY if action == "finalize" else document.Status.REVIEW
                document.save(update_fields=["title", "edited_data", "status", "error_message", "disk_result_key", "updated_at"])
                messages.success(
                    request,
                    "Документ сохранён и готов к скачиванию." if action == "finalize" else
                    "Предпросмотр обновлён." if action == "regenerate" else "Изменения сохранены.",
                )
                return redirect("studio:review", pk=document.pk)
            except Exception as exc:
                document.status = document.Status.ERROR
                document.error_message = str(exc)
                document.save(update_fields=["status", "error_message", "updated_at"])
                messages.error(request, f"Не удалось сгенерировать документ: {exc}")
    else:
        form = DocumentReviewForm(document=document)
    return render(request, "studio/review.html", {"document": document, "form": form})


@login_required
@xframe_options_sameorigin
def preview_document(request, pk):
    """Lightweight local preview when LibreOffice PDF conversion is unavailable."""
    from html import escape
    from docx import Document

    document = get_object_or_404(SourceDocument, pk=pk)
    if not document.generated_docx:
        return HttpResponse("Документ ещё не сгенерирован.", status=404, content_type="text/plain; charset=utf-8")
    word = Document(document.generated_docx.path)

    def paragraph_html(paragraph):
        text = escape(paragraph.text).replace("\n", "<br>")
        style = (paragraph.style.name if paragraph.style else "").lower()
        tag = "h1" if "title" in style or "heading 1" in style else "h2" if "heading 2" in style else "p"
        align = {1: "center", 2: "right", 3: "justify"}.get(paragraph.alignment, "left")
        return f'<{tag} style="text-align:{align}">{text or "&nbsp;"}</{tag}>'

    pieces = []
    for section in word.sections:
        header = "".join(paragraph_html(p) for p in section.header.paragraphs if p.text.strip())
        if header:
            pieces.append(f'<header class="word-header">{header}</header>')
    pieces.extend(paragraph_html(p) for p in word.paragraphs)
    for table in word.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{escape(cell.text).replace(chr(10), '<br>')}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        pieces.append(f'<table>{"".join(rows)}</table>')
    for section in word.sections:
        footer = "".join(paragraph_html(p) for p in section.footer.paragraphs if p.text.strip())
        if footer:
            pieces.append(f'<footer class="word-footer">{footer}</footer>')

    html = """<!doctype html><html lang="ru"><head><meta charset="utf-8"><style>
      *{box-sizing:border-box}body{margin:0;padding:36px;background:#ecedf2;color:#111;font:15px/1.55 Arial,sans-serif}
      .page{width:min(820px,100%);min-height:1060px;margin:auto;padding:72px 78px;background:white;box-shadow:0 12px 35px rgba(49,38,59,.14)}
      p{margin:0 0 11px;white-space:pre-wrap}h1{font-size:24px;margin:0 0 22px}h2{font-size:19px;margin:20px 0 10px}
      table{width:100%;margin:14px 0 20px;border-collapse:collapse}td{padding:8px 10px;border:1px solid #aaa;vertical-align:top}
      .word-header{margin-bottom:26px;padding-bottom:10px;border-bottom:1px solid #ddd;color:#666;font-size:12px}
      .word-footer{margin-top:35px;padding-top:10px;border-top:1px solid #ddd;color:#666;font-size:12px}
      @media(max-width:700px){body{padding:10px}.page{padding:35px 25px;min-height:90vh}}
    </style></head><body><article class="page">""" + "".join(pieces) + "</article></body></html>"
    return HttpResponse(html)


@login_required
@require_POST
def reanalyze_document(request, pk):
    document = get_object_or_404(SourceDocument, pk=pk)
    try:
        process_document(document)
        messages.success(request, "ИИ повторно проанализировал скан.")
    except Exception as exc:
        messages.error(request, f"Повторный анализ не выполнен: {exc}")
    return redirect("studio:review", pk=document.pk)


@login_required
@xframe_options_sameorigin
def document_file(request, pk, field_name):
    """Serve private scans/results only to authenticated users."""
    if field_name not in {"source_pdf", "generated_docx", "generated_pdf"}:
        raise Http404
    document = get_object_or_404(SourceDocument, pk=pk)
    stored_file = getattr(document, field_name)
    if not stored_file:
        raise Http404
    path = Path(stored_file.path)
    if not path.is_file():
        raise Http404
    content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    return FileResponse(
        path.open("rb"),
        as_attachment=request.GET.get("download") == "1",
        filename=path.name,
        content_type=content_type,
    )
