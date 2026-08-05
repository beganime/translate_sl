"""Create completed working templates while preserving all supplied originals."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]


def apply_font(run, name="Times New Roman", size=11, bold=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def replace_paragraph(paragraph, text, *, size=11, bold=None, alignment=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    apply_font(run, size=size, bold=bold)
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def complete_foreign_passport():
    source = ROOT / "Загран_доделанный.docx"
    output = ROOT / "Загран_доделанный_v2.docx"
    doc = Document(source)

    labels = replace_paragraph(doc.paragraphs[14], "ПАСПОРТ\tВид:\tНазвание страны:\tНомер паспорта:", size=10)
    values = replace_paragraph(doc.paragraphs[15], "\t{{ view }}\tТКМ\t{{ passport }}", size=12, bold=True)
    for paragraph in (labels, values):
        stops = paragraph.paragraph_format.tab_stops
        stops.clear_all()
        stops.add_tab_stop(Inches(2.65), WD_TAB_ALIGNMENT.CENTER)
        stops.add_tab_stop(Inches(4.55), WD_TAB_ALIGNMENT.CENTER)
        stops.add_tab_stop(Inches(6.55), WD_TAB_ALIGNMENT.CENTER)

    replace_paragraph(
        doc.paragraphs[30],
        "Р<TKM{{ surname | upper }}<<{{ name | upper }}<<<<<<<<<<<<<<",
        size=12,
        bold=True,
    )
    replace_paragraph(
        doc.paragraphs[31],
        "{{ passport_num | upper }}<{{ passport_code | upper }}<<<<{{ alike | upper }}",
        size=12,
        bold=True,
    )
    doc.save(output)
    print(output)


def complete_birth_certificate():
    source = ROOT / "свид_о_рож_не_доделанный.docx"
    output = ROOT / "Свидетельство_о_рождении_доделанный.docx"
    doc = Document(source)
    replacements = {
        2: "{{ certificate_number }}",
        4: "Гражданин(ка): {{ child_full_name }}",
        5: "Дата рождения: {{ birth_date }}",
        6: "{{ birth_date_words }}",
        7: "Место рождения: {{ birth_place }}",
        8: "Дата регистрации акта о рождении: {{ registration_date }}",
        9: "Произведена запись за № {{ record_number }}",
        13: "Отец: {{ father_full_name }}",
        14: "Национальность: {{ father_nationality }}",
        16: "Мать: {{ mother_full_name }}",
        17: "Национальность: {{ mother_nationality }}",
        21: "Место регистрации: {{ registration_office }}",
        23: "Дата выдачи: {{ issue_date }}",
        28: "{{ seal_text }}",
    }
    for index, text in replacements.items():
        replace_paragraph(doc.paragraphs[index], text, size=11)
    doc.save(output)
    print(output)


def complete_internal_passport():
    source = ROOT / "внутренний_паспорт_не_доделанный.docx"
    output = ROOT / "Внутренний_паспорт_доделанный.docx"
    doc = Document(source)
    replacements = {
        2: "{{ full_name }}",
        3: "{{ series_number }}",
        5: "Дата рождения: {{ birth_date }}",
        6: "Место рождения: {{ birth_place }}",
        7: "Национальность: {{ nationality }}",
        8: "Кем выдан паспорт: {{ issuer }}",
        10: "{{ issue_date }}",
        11: "{{ seal_text }}",
        38: "{{ mrz_line_1 | upper }}",
        39: "{{ mrz_line_2 | upper }}",
    }
    for index, text in replacements.items():
        replace_paragraph(doc.paragraphs[index], text, size=11)

    aligned_rows = {
        27: "\tКод страны:\tНомер паспорта:",
        28: "\t{{ country_code }}\t{{ passport_number }}",
        29: "\tФамилия:",
        30: "\t{{ surname | upper }}",
        31: "ФОТО\tИмя:",
        32: "ВЛАДЕЛЬЦА\t{{ name | upper }}",
        33: "\tДата рождения:\tМесто рождения:",
        34: "\t{{ birth_date }}\t{{ birth_place_short }}",
        35: "\tПол:\tДата выдачи:",
        36: "\t{{ sex }}\t{{ issue_date }}",
    }
    for index, text in aligned_rows.items():
        paragraph = replace_paragraph(doc.paragraphs[index], text, size=10, bold=index in {27, 29, 31, 33, 35})
        stops = paragraph.paragraph_format.tab_stops
        stops.clear_all()
        stops.add_tab_stop(Inches(2.75), WD_TAB_ALIGNMENT.CENTER)
        stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.CENTER)

    for index, row in enumerate(doc.tables[0].rows, start=1):
        cell = row.cells[0]
        cell.text = (
            f"{{{{ registration_authority_{index} }}}}\n"
            f"ПРОПИСАН по адресу: {{{{ registration_address_{index} }}}}\n"
            f"{{{{ registration_date_{index} }}}}\n"
            "Подпись: /подпись имеется/"
        )
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                apply_font(run, size=10)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    complete_foreign_passport()
    complete_birth_certificate()
    complete_internal_passport()
