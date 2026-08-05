"""Complete the supplied attestation template without overwriting the source."""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Аттестат_доделанный.docx"
ORIGINAL = ROOT / "Аттестат_не_доделанный_шаблон.docx"
SOURCE = ORIGINAL if ORIGINAL.exists() else OUTPUT

SUBJECTS = [
    ("Туркменский язык", "{{ turkmen_lang }}"),
    ("Туркменская литература", "{{ turkmen_literature }}"),
    ("Родной язык", "{{ russuian_lang }}"),
    ("Русская литература", "{{ russian_literature }}"),
    ("Иностранный язык: английский язык\n                              русский язык", "{{ english_lang }}\n{{ russian_lang }}"),
    ("Алгебра и начала анализа", "{{ algebra }}"),
    ("Геометрия", "{{ geometry }}"),
    ("Информатика", "{{ informatics }}"),
    ("Информационно-коммуникационные и инновационные технологии", "{{ communications }}"),
    ("Физика", "{{ physics }}"),
    ("Астрономия", "{{ astronomy }}"),
    ("Химия", "{{ chemistry }}"),
    ("Биология", "{{ biology }}"),
    ("Экология", "{{ ecology }}"),
    ("География", "{{ geography }}"),
    ("История Туркменистана", "{{ history_turkmenistan }}"),
    ("Всеобщая история", "{{ world_history }}"),
    ("Обществоведение", "{{ social_science }}"),
    ("Основы государства и права Туркменистана", "{{ state_and_law }}"),
    ("Основы экономики", "{{ economics }}"),
    ("Основы современных технологий", "{{ modern_technologies }}"),
    ("Моделирование и графика", "{{ modeling_and_graphics }}"),
    ("Культура поведения", "{{ ethics }}"),
    ("Культурное наследие Туркменистана", "{{ cultural_heritage }}"),
    ("Мировая культура", "{{ world_culture }}"),
    ("Основы проектирования", "{{ project_design }}"),
    ("Физкультура", "{{ physical_education }}"),
    ("Основы безопасности жизнедеятельности", "{{ life_safety }}"),
]


def set_cell(cell, text, *, bold=False, center=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        if center:
            paragraph.alignment = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(8.5)
            run.bold = bold


def main():
    doc = Document(SOURCE)
    intro = doc.paragraphs[5]
    intro.clear()
    intro_run = intro.add_run(
        "Настоящий аттестат выдан: {{ fio }}. Дата рождения: {{ date_of_birth }}. "
        "Место рождения: {{ city }}. В {{ year }} году завершено обучение в {{ school_name }} "
        "({{ school_address }}). За время обучения получены следующие оценки:"
    )
    intro_run.font.name = "Times New Roman"
    intro_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    intro_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    intro_run.font.size = Pt(11)
    intro.alignment = 3
    intro.paragraph_format.space_after = Pt(4)
    table = doc.tables[0]
    behavior_row = table.rows[-1]._tr

    # The source has 27 subject slots (one blank); the scan has 28 subjects.
    if len(table.rows) < len(SUBJECTS) + 2:
        prototype = deepcopy(table.rows[-2]._tr)
        behavior_row.addprevious(prototype)

    set_cell(table.rows[0].cells[0], "Наименование предмета", bold=True, center=True)
    set_cell(table.rows[0].cells[1], "Оценка", bold=True, center=True)
    for index, (label, variable) in enumerate(SUBJECTS, start=1):
        set_cell(table.rows[index].cells[0], label)
        set_cell(table.rows[index].cells[1], variable, center=True)

    # Preserve the source behaviour row but ensure its Jinja slot is clean.
    final_row = table.rows[len(SUBJECTS) + 1]
    merged = final_row.cells[0].merge(final_row.cells[1])
    set_cell(merged, "Поведение: {{ behavior }}")

    seal_paragraph = next(p for p in doc.paragraphs if p.text.strip().startswith("/Печать:"))
    seal_paragraph.clear()
    seal_paragraph.alignment = 3
    label = seal_paragraph.add_run("/Печать:")
    label.bold = True
    label.italic = True
    value = seal_paragraph.add_run(" {{ seal_text }}")
    value.italic = True
    for run in seal_paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        run.font.size = Pt(11)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
