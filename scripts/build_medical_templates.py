"""Build Jinja-ready medical DOCX templates from the four retained references."""

from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate


ROOT = Path(__file__).resolve().parents[1]


SOURCES = {
    "spid": (
        ROOT / "спид_недоделанный.docx",
        ROOT / "медсправка_ВИЧ_шаблон.docx",
        "5cdeafb9b3909eebb5784d4d71882c159b4141320c9f64106c728e3d8d002314",
    ),
    "hepatitis": (
        ROOT / "геппатит_недоделанный.docx",
        ROOT / "медсправка_гепатит_шаблон.docx",
        "899b0d05f58133d67476705c729e0a4b02c309e176dda5365758bb32c24e044c",
    ),
    "form086": (
        ROOT / "форма_086_недоделанный.docx",
        ROOT / "медсправка_086_шаблон.docx",
        "ece9d1d2f3a28f28006a8d874b53ca0c54a4ab1401f21a133e0eb24c50ee545e",
    ),
    "tb": (
        ROOT / "тубмедсправка_недоделанный.docx",
        ROOT / "медсправка_туберкулёз_шаблон.docx",
        "6e9dce038ecf781d8d8ac177e69fd9b066ed818b8bc7c8d662594e49782fee2f",
    ),
}


EXPECTED_VARIABLES = {
    "spid": {
        "medical_organization", "certificate_number", "patient_full_name",
        "examination_result", "citizenship", "destination_country", "birth_date",
        "passport_number", "validity_period", "issue_date", "doctor_name",
        "commission_chair_name", "seal_text",
    },
    "hepatitis": {
        "medical_organization", "report_number", "patient_full_name", "birth_date",
        "sex", "address", "additional_information", "hbsag_cutoff_od",
        "hbsag_sample_od", "hbsag_result", "hbsag_report_number", "hbsag_test_date",
        "anti_hcv_cutoff_od", "anti_hcv_sample_od", "anti_hcv_result",
        "anti_hcv_report_number", "anti_hcv_test_date", "result_issue_date",
        "performer", "department_head", "seal_text",
    },
    "form086": {
        "stamp_health_authority", "stamp_medical_organization", "stamp_number",
        "stamp_date", "medical_organization", "certificate_number",
        "consultation_date", "issuer", "destination_institution", "patient_full_name",
        "sex", "birth_date", "address", "past_illnesses", "therapist",
        "surgeon", "neurologist", "ophthalmologist", "otolaryngologist",
        "other_specialists", "narcologist", "psychiatrist", "dermatovenerologist",
        "xray_result", "laboratory_result", "vaccinations", "commission_chair",
        "institution_head", "seal_text",
    },
    "tb": {
        "medical_organization", "stamp_number", "stamp_date", "patient_full_name_dative",
        "birth_year", "xray_date", "issue_date", "destination_country",
        "chief_physician", "seal_text",
    },
}


def paragraphs(document):
    """Yield unique body/table/header/footer paragraphs."""
    seen = set()

    def emit(items):
        for paragraph in items:
            marker = paragraph._p
            if marker not in seen:
                seen.add(marker)
                yield paragraph

    yield from emit(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from emit(cell.paragraphs)
    for section in document.sections:
        yield from emit(section.header.paragraphs)
        yield from emit(section.footer.paragraphs)


def replace_in_runs(paragraph, old: str, new: str) -> None:
    """Replace one substring while retaining the formatting of its first run."""
    full = "".join(run.text for run in paragraph.runs)
    start = full.find(old)
    if start < 0:
        raise ValueError(f"Не найден текст {old!r} в абзаце {full!r}")
    if full.find(old, start + len(old)) >= 0:
        raise ValueError(f"Текст {old!r} встречается в абзаце несколько раз")

    end = start + len(old)
    offsets = []
    cursor = 0
    for run in paragraph.runs:
        offsets.append((cursor, cursor + len(run.text)))
        cursor += len(run.text)
    first = next(i for i, (_, run_end) in enumerate(offsets) if run_end > start)
    last = next(i for i, (run_start, _) in enumerate(offsets) if run_start < end <= offsets[i][1])
    first_start, _ = offsets[first]
    last_start, _ = offsets[last]
    prefix = paragraph.runs[first].text[: start - first_start]
    suffix = paragraph.runs[last].text[end - last_start :]
    if first == last:
        paragraph.runs[first].text = prefix + new + suffix
    else:
        paragraph.runs[first].text = prefix + new
        for index in range(first + 1, last):
            paragraph.runs[index].text = ""
        paragraph.runs[last].text = suffix


def replace(document, marker: str, old: str, new: str) -> None:
    candidates = [p for p in paragraphs(document) if marker in p.text and old in p.text]
    if len(candidates) != 1:
        raise ValueError(f"Ожидался один абзац marker={marker!r}, old={old!r}; найдено {len(candidates)}")
    replace_in_runs(candidates[0], old, new)


def clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def set_cell_value(cell, value: str) -> None:
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = value
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(value)
    for paragraph in cell.paragraphs[1:]:
        clear_paragraph(paragraph)


def set_paragraph_value(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def build_spid(source: Path, output: Path) -> None:
    doc = Document(source)
    set_paragraph_value(doc.tables[0].cell(0, 0).paragraphs[4], "{{ medical_organization }}")
    replace(doc, "Медицинский сертификат", "939498", "{{ certificate_number }}")
    replace(doc, "Выдан:", "Выдан:", "Выдано лицу:")
    replace(doc, "Выдано лицу:", "Рахманова Махри", "{{ patient_full_name }}")
    replace(doc, "Выдано лицу:", ", который прошол медицинское обследование в", ". Медицинское обследование проведено медицинской организацией:")
    replace(doc, "Центре профилактики", "Центре профилактики заболеваний СПИДом Лебапской области ", "{{ medical_organization }}")
    replace(doc, "Анализ на наличие", "Анализ на наличие антител к ВИЧ с отрицательным результатом.", "{{ examination_result }}")
    replace(doc, "Гражданину", "Гражданину  по состоянию здоровья:  ", "Гражданство: ")
    replace(doc, "Гражданство:", "из Туркменистана ", "{{ citizenship }}")
    replace(doc, "По  месту", "Россия", "{{ destination_country }}")
    replace(doc, "Год рождения", "04.06.2005", "{{ birth_date }}")
    replace(doc, "Номер паспорта", "А 2120872", "{{ passport_number }}")
    replace(doc, "Срок действия", "3 месяца", "{{ validity_period }}")
    replace(doc, "Дата выдачи", "10.03.2026", "{{ issue_date }}")
    replace(doc, "Врач выдавший", "Розыева М.Г", "{{ doctor_name }}")
    replace(doc, "Председатель комиссии", "Нуруев Б. Н.", "{{ commission_chair_name }}")
    replace(doc, "Перевод печати", "Министерство  здравоохранения  и  медицинской  промышленности  Туркменистана,", "{{ seal_text }}")
    seal_tail = next(p for p in doc.paragraphs if "Центр  по  профилактике" in p.text)
    clear_paragraph(seal_tail)
    doc.save(output)


def build_hepatitis(source: Path, output: Path) -> None:
    doc = Document(source)
    replace(doc, "Наименование организации", "Служба санитарного контроля и контроля заболеваний Лебапской области", "{{ medical_organization }}")
    replace(doc, "Результаты иммуноферментного", "2910", "{{ report_number }}")
    replace(doc, "Ф. И.О.", "Астанова Гулалла", "{{ patient_full_name }}")
    replace(doc, "Дата рождение", "Дата рождение", "Дата рождения")
    replace(doc, "Дата рождения", "28/01/2008", "{{ birth_date }}")
    replace(doc, "Пол:", "Ж", "{{ sex }}")
    replace(doc, "Адрес:", "г.Туркменабат,  ул. Бахар дом 16", "{{ address }}")
    replace(doc, "Дополнительная информация", "65868322", "{{ additional_information }}")
    values = [
        (1, 1, 1, "{{ hbsag_cutoff_od }}"), (1, 2, 1, "{{ hbsag_sample_od }}"),
        (1, 3, 1, "{{ hbsag_result }}"), (1, 4, 1, "{{ hbsag_report_number }}"),
        (1, 5, 1, "{{ hbsag_test_date }}"), (2, 1, 1, "{{ anti_hcv_cutoff_od }}"),
        (2, 2, 1, "{{ anti_hcv_sample_od }}"), (2, 3, 1, "{{ anti_hcv_result }}"),
        (2, 4, 1, "{{ anti_hcv_report_number }}"), (2, 5, 1, "{{ anti_hcv_test_date }}"),
    ]
    for table_index, row, column, value in values:
        set_cell_value(doc.tables[table_index].cell(row, column), value)
    replace(doc, "Дата выдачи результата", "25.06.2026", "{{ result_issue_date }}")
    replace(doc, "Исполнитель", "Хаитова Ф.М.", "{{ performer }}")
    replace(doc, "Задующий отделом", "Задующий", "Заведующий")
    replace(doc, "Заведующий отделом", "Хаитова Ф.М.", "{{ department_head }}")
    seal_cell = doc.tables[3].cell(0, 0)
    replace_in_runs(seal_cell.paragraphs[0], "Министерство Здравоохранения и", "{{ seal_text }}")
    for paragraph in seal_cell.paragraphs[1:]:
        clear_paragraph(paragraph)
    doc.save(output)


def build_form086(source: Path, output: Path) -> None:
    doc = Document(source)
    replace_in_runs(doc.paragraphs[3], "Управление здравоохранения Лебапской области", "{{ stamp_health_authority }}")
    replace_in_runs(doc.paragraphs[4], "Туркменабатская городская поликлиника № 2", "{{ stamp_medical_organization }}")
    set_paragraph_value(
        doc.paragraphs[5],
        "{% if stamp_number or stamp_date %}№ {{ stamp_number }}\t\t\t{{ stamp_date }} г.{% endif %}",
    )
    replace(doc, "Наименование учреждения", "-", "{{ medical_organization }}")
    replace(doc, "Медицинская справка", "37", "{{ certificate_number }}")
    replace(doc, "поступающих на работу", "10 »  03     2026", "{{ consultation_date }}")
    replace(doc, "1. Кем выдано", "№ 2 дом здоровье города Туркменабат", "{{ issuer }}")
    replace(doc, "2. Наименование", "-", "{{ destination_institution }}")
    replace(doc, "3.Фамилия", "Рахманова Махри Шохрадовна", "{{ patient_full_name }}")
    replace(doc, "4. Пол", "женский", "{{ sex }}")
    replace(doc, "5. Дата рождения", "04.06.2005", "{{ birth_date }}")
    replace(doc, "6. Адрес", "город Туркменабат Лебапской области улица 2-ой жилой комплекс дом 15 кв 48", "{{ address }}")
    replace(doc, "7. Перенесённые", "нет", "{{ past_illnesses }}")
    specialist_values = [
        ("Терапевт:", "Здорова (подпись и печать врача имеется) 13.03.2026 г.", "{{ therapist }}"),
        ("Хирург:", "Здорова (подпись и печать врача имеется) 13.03.2026 г.", "{{ surgeon }}"),
        ("Невропатолог:", "Здорова (подпись и печать врача имеется) 13.03.2026 г.", "{{ neurologist }}"),
        ("Окулист:", "Зрение 1,0 здорова (подпись и печать врача имеется) 13.03.2026 г.", "{{ ophthalmologist }}"),
        ("Отоларинголог:", "Здорова (подпись и печать врача имеется) 13.03.2026 г.", "{{ otolaryngologist }}"),
        ("Нарколог:", "На учёте не состоит. (подпись врача имеется) 13.03.2026 г. № 4101", "{{ narcologist }}"),
        ("Психиатр:", "На учёте не состоит. (подпись врача имеется) 13.03.2026 г.", "{{ psychiatrist }}"),
        ("Кож вен:", "На учёте не состоит. (подпись врача имеется) 13.03.2026 г. № 791", "{{ dermatovenerologist }}"),
    ]
    for marker, old, new in specialist_values:
        replace(doc, marker, old, new)
    replace(doc, "Другие специалисты", "Другие специалисты: ", "Другие специалисты: {{ other_specialists }}")
    replace(doc, "9. Данные рентгенологические", "На флюорограмме сердце и легкие в норме (подпись врача имеется) 13.03.2026 г.", "{{ xray_result }}")
    replace(doc, "10. Данные лабораторных", "Общий анализ крови и мочи в норме. (подпись врача имеется)", "{{ laboratory_result }}")
    replace(doc, "11. Профилактические", "по возрасту годна к учёбе", "{{ vaccinations }}")
    replace(doc, "Председатель комиссии", "Караева Р.Т", "{{ commission_chair }}")
    replace(doc, "Руководитель лечебно", "Джуманазарова Р.Н.", "{{ institution_head }}")
    seal_cell = doc.tables[1].cell(0, 0)
    replace_in_runs(seal_cell.paragraphs[0], "МЗ и МП Туркменистана", "{{ seal_text }}")
    for paragraph in seal_cell.paragraphs[1:]:
        clear_paragraph(paragraph)
    doc.save(output)


def build_tb(source: Path, output: Path) -> None:
    doc = Document(source)
    replace_in_runs(doc.paragraphs[1], "Лебапская областная противотуберкулёзная больница", "{{ medical_organization }}")
    replace(doc, "№ 822", "822", "{{ stamp_number }}")
    replace(doc, "№ {{ stamp_number }}", "13.03.2026", "{{ stamp_date }}")
    replace(doc, "Выдана Рахмановой", "Рахмановой Махри", "{{ patient_full_name_dative }}")
    replace(doc, "года рождения", "2005", "{{ birth_year }}")
    replace(doc, "она действительно", "она действительно", "Указанное лицо действительно")
    replace(doc, "рентген-флюорографическом", "13.03.2026", "{{ xray_date }}")
    replace(doc, "Дата:", "13.03.2026", "{{ issue_date }}")
    replace(doc, "месту требованию", "Россия", "{{ destination_country }}")
    replace(doc, "Главный врач", "Ходжамбердиев  Д.Р.", "{{ chief_physician }}")
    seal = next(p for p in doc.paragraphs if p.text.startswith("Перевод печати"))
    seal.runs[0].text = "Перевод печати: "
    seal.runs[1].text = "{{ seal_text }}"
    for run in seal.runs[2:]:
        run.text = ""
    doc.save(output)


BUILDERS = {
    "spid": build_spid,
    "hepatitis": build_hepatitis,
    "form086": build_form086,
    "tb": build_tb,
}


def main() -> None:
    for slug, (source, output, expected_hash) in SOURCES.items():
        actual_hash = hashlib.sha256(source.read_bytes()).hexdigest()
        if actual_hash != expected_hash:
            raise RuntimeError(f"Эталон {source.name} изменён: {actual_hash} != {expected_hash}")
        BUILDERS[slug](source, output)
        found = set(DocxTemplate(output).get_undeclared_template_variables())
        if found != EXPECTED_VARIABLES[slug]:
            raise RuntimeError(f"Неверный набор переменных {slug}: missing={EXPECTED_VARIABLES[slug]-found}, extra={found-EXPECTED_VARIABLES[slug]}")
        print(f"Создан {output.name}: {len(found)} переменных")


if __name__ == "__main__":
    main()
