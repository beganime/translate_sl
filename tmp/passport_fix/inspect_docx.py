import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


def paragraph_info(paragraph, path):
    return {
        "path": path,
        "text": paragraph.text,
        "alignment": str(paragraph.alignment),
        "style": paragraph.style.name if paragraph.style else None,
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "size_pt": run.font.size.pt if run.font.size else None,
                "font": run.font.name,
            }
            for run in paragraph.runs
        ],
    }


def inspect(path):
    doc = Document(path)
    paragraphs = [paragraph_info(p, f"body.p[{i}]") for i, p in enumerate(doc.paragraphs)]
    tables = []
    for ti, table in enumerate(doc.tables):
        cells = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cells.append({
                    "path": f"table[{ti}].row[{ri}].cell[{ci}]",
                    "width": cell.width,
                    "paragraphs": [paragraph_info(p, f"p[{pi}]") for pi, p in enumerate(cell.paragraphs)],
                })
        tables.append({"rows": len(table.rows), "cols": len(table.columns), "cells": cells})
    with zipfile.ZipFile(path) as package:
        xml = etree.fromstring(package.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main", "v": "urn:schemas-microsoft-com:vml", "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape", "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
        shapes = {
            "vml_rect": len(xml.xpath(".//v:rect", namespaces=ns)),
            "vml_shape": len(xml.xpath(".//v:shape", namespaces=ns)),
            "word_shapes": len(xml.xpath(".//wps:wsp", namespaces=ns)),
            "drawings": len(xml.xpath(".//wp:anchor | .//wp:inline", namespaces=ns)),
        }
        parts = sorted(package.namelist())
    return {"file": str(path), "paragraphs": paragraphs, "tables": tables, "shapes": shapes, "parts": parts}


for filename in sys.argv[1:]:
    path = Path(filename)
    output = path.with_suffix(".structure.json")
    output.write_text(json.dumps(inspect(path), ensure_ascii=False, indent=2), encoding="utf-8")
    print(output)
